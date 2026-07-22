from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document


def run_details(run) -> dict[str, object]:
    color = run.font.color.rgb
    return {
        "text": run.text,
        "bold": bool(run.bold),
        "italic": bool(run.italic),
        "underline": bool(run.underline),
        "highlight": str(run.font.highlight_color) if run.font.highlight_color else None,
        "color": str(color) if color else None,
    }


def inspect(path: Path) -> dict[str, object]:
    document = Document(path)
    return {
        "path": str(path),
        "paragraphs": [
            {
                "index": index,
                "style": paragraph.style.name,
                "text": paragraph.text,
                "runs": [run_details(run) for run in paragraph.runs if run.text],
            }
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ],
        "tables": [
            [
                [
                    {
                        "text": cell.text,
                        "paragraphs": [
                            {
                                "text": paragraph.text,
                                "runs": [run_details(run) for run in paragraph.runs if run.text],
                            }
                            for paragraph in cell.paragraphs
                        ],
                    }
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in document.tables
        ],
    }


for raw_path in sys.argv[1:]:
    print(json.dumps(inspect(Path(raw_path)), ensure_ascii=False, indent=2))
